import os
import re
import docx
import unicodedata
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Cm
from nltk import word_tokenize, sent_tokenize
import PyPDF2
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from .helper import archivos_referencia_limpios
from .metodos_de_similitud import obtener_similitud


class ArchivoTxt:
    def __init__(self, nombre, extension, texto):
        self.nombre = nombre
        self.extension = extension
        self.texto = texto


def obtener_archivos(nombre_directorio, lista_archivos_adicionales=[]):
    archivos = os.listdir(nombre_directorio)
    # print("nombre_directorio =",nombre_directorio)
    # print(archivos)
    archivos_txt = [convertir_archivo_a_txt(nombre_directorio, archivo) for archivo in archivos]
    if lista_archivos_adicionales:
        for archivo_adicional in lista_archivos_adicionales:
            archivo_txt = convertir_archivo_adicional_a_txt(archivo_adicional)
            archivos_txt.append(archivo_txt)
    
    return archivos_txt

def obtener_archivo_Test(ruta_completa):
    directorio, nombre_archivo = os.path.split(ruta_completa)
    return [convertir_archivo_a_txt(directorio+"/", nombre_archivo)]

#cambio
def convertir_archivo_a_txt(nombre_directorio, archivo):
    archivo_nombre, archivo_extension = os.path.splitext(archivo)
    archivo_completo = os.path.join(nombre_directorio, archivo)
    try:
        with open(archivo_completo, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            archivo_txt = '\n'.join(page.extract_text() for page in pdf_reader.pages)
    except Exception as e:
        archivo_txt = ''  # Manejar el error de lectura de PDF vacío
    return ArchivoTxt(archivo_nombre, archivo_extension, archivo_txt)

def convertir_archivo_adicional_a_txt(path):
    nombre_directorio, archivo = os.path.split(path)
    archivo_nombre, archivo_extension = os.path.splitext(archivo)
    archivo_completo = os.path.join(nombre_directorio, archivo)
    try:
        with open(archivo_completo, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            archivo_txt = '\n'.join(page.extract_text() for page in pdf_reader.pages)
    except Exception as e:
        archivo_txt = ''  # Manejar el error de lectura de PDF vacío
    return ArchivoTxt(archivo_nombre, archivo_extension, archivo_txt)

def limpieza(archivo):
    archivo_limpio = re.sub(r'\n+', '\n', archivo.strip())
    archivo_limpio = re.sub('\n', '. ', archivo_limpio.strip())
    archivo_limpio = re.sub(r'[.][.]+', '.', archivo_limpio.strip())
    archivo_limpio = re.sub(r'[ ][ ]+', ' ', archivo_limpio.strip())
    archivo_limpio = re.sub('á', 'a', archivo_limpio.strip())
    archivo_limpio = re.sub('é', 'e', archivo_limpio.strip())
    archivo_limpio = re.sub('í', 'i', archivo_limpio.strip())
    archivo_limpio = re.sub('ó', 'o', archivo_limpio.strip())
    archivo_limpio = re.sub('ú', 'u', archivo_limpio.strip())
    archivo_limpio = re.sub('”', '"', archivo_limpio.strip())
    archivo_limpio = re.sub('“', '"', archivo_limpio.strip())
    archivo_limpio = re.sub('\u200b', ' ', archivo_limpio.strip())

    archivo_limpio = unicodedata.normalize("NFKD", archivo_limpio.strip())

    oraciones = sent_tokenize(archivo_limpio.strip(), "spanish")
    oraciones_limpias = []
    for oracion in oraciones:
        if oracion.strip() != '.':
            if oracion.strip().endswith('.'):
                oracion_a_agregar = oracion[:-1]
            else:
                oracion_a_agregar = oracion
            oraciones_limpias.append(oracion_a_agregar.strip())

    i=0
    j=0

    oraciones_mas_limpias = []
    while i < len(oraciones_limpias):
        if i == 0:
            oraciones_mas_limpias.append(oraciones_limpias[0])
        else:
            palabras_oracion = word_tokenize(oraciones_limpias[i])
            if palabras_oracion[0].islower():
                oraciones_mas_limpias[j] += " " + oraciones_limpias[i]
            else:
                j += 1
                oraciones_mas_limpias.append(oraciones_limpias[i])
        i += 1
    return oraciones_mas_limpias


def limpiar_archivos_referencia(archivo):
    archivo_limpio = limpieza(archivo.texto)
    archivos_referencia_limpios.append(ArchivoTxt(archivo.nombre, archivo.extension, archivo_limpio))


def correctamente_citada(url, texto_archivo_test_limpio):
    for oracion in texto_archivo_test_limpio:
        similitud = obtener_similitud(oracion, url)
        if similitud >= 0.99:
            return True
    return False


def excluida(oracion,path): #aqui
    archivo_text_excluido = convertir_archivo_a_txt(path, "Texto excluido de plagio.txt")
    if archivo_text_excluido.texto is None:
        return False
    archivo_limpio = limpieza(archivo_text_excluido.texto)
    for oracion_archivo in archivo_limpio:
        similitud = obtener_similitud(oracion, oracion_archivo)
        if similitud > 0.9:
            return True
    return False


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run ()
    r._r.append (hyperlink)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    return hyperlink


# def guardar_resultado(nombre_archivo,  topico_con_mas_score, plagio, tiempo_que_tardo, porcentaje_de_plagio, path_resultado, path_referencia): #aqui
#     # Crear documento Word con los resultados y guardarlo como un nuevo archivo
#     document = Document()

#     h = document.add_heading(f'Análisis de plagio sobre:\n', 0)
#     h.add_run(nombre_archivo).italic = True

#     p = document.add_paragraph('Tópicos del texto: ')
#     p.add_run(", ".join(topico_con_mas_score)).italic = True

#     document.add_heading('Análisis de plagio', level=1)
#     document.add_paragraph(f'Total de {len(plagio)} plagios encontrados en {tiempo_que_tardo}')
#     document.add_paragraph(f'Porcentaje de plagio general: {porcentaje_de_plagio}%')

#     table = document.add_table(rows=1, cols=4)
#     table.style = 'Medium Shading 1 Accent 1'
#     table.autofit = False

#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'Oración plagiada'
#     hdr_cells[1].text = 'Oración original'
#     hdr_cells[2].text = 'Lugar donde se encontró'
#     hdr_cells[3].text = 'Ubicación'

#     for oracion, plagio, porcentaje, url, ubicacion in plagio:
#         row_cells = table.add_row().cells
#         row_cells[0].text = oracion
#         row_cells[1].text = plagio
#         p = row_cells[2].add_paragraph()
#         if str(url).startswith("http"):
#             add_hyperlink(p, url, url)
#         else:
#             add_hyperlink(p, url, os.path.abspath(path_referencia) + '\\' + url)
#         row_cells[3].text = ubicacion

#     widths = (Cm(5), Cm(5), Cm(3.47), Cm(2.15))
#     for row in table.rows:
#         for idx, width in enumerate(widths):
#             row.cells[idx].width = width

#     # Guardar el documento Word
#     nombre_archivo_plagio = path_resultado + 'Plagio ' + str(str(nombre_archivo).split(".")[0]) + '.docx'
#     document.save(nombre_archivo_plagio)
#     print("nombre_archivo_plagio",nombre_archivo_plagio)
#     # Convertir el documento Word a PDF
#     nombre_archivo_pdf = path_resultado + 'Plagio ' + str(str(nombre_archivo).split(".")[0]) + '.pdf'
#     convert_to_pdf(nombre_archivo_plagio, nombre_archivo_pdf) #quitar
#     os.remove(nombre_archivo_plagio)
#     nombre ='Plagio ' + str(str(nombre_archivo).split(".")[0]) + '.pdf'
#     return nombre_archivo_pdf, nombre

def guardar_resultado(nombre_archivo, plagio, tiempo_que_tardo, porcentaje_de_plagio, path_resultado, path_referencia):
    # Crear el contenido del PDF
    pdf_content = []

    # Agregar el título y subtítulo al PDF
    pdf_content.append(Paragraph(f'Análisis de plagio sobre:\n', getSampleStyleSheet().get('Heading1')))
    pdf_content.append(Paragraph(nombre_archivo, getSampleStyleSheet().get('Title')))
    pdf_content.append(Spacer(1, 12))  # Espacio en blanco entre párrafos

    # Agregar información del análisis de plagio al PDF
    pdf_content.append(Paragraph(f'Análisis de plagio', getSampleStyleSheet().get('Heading1')))
    pdf_content.append(Paragraph(f'Total de {len(plagio)} plagios encontrados en {tiempo_que_tardo}', getSampleStyleSheet().get('Normal')))
    pdf_content.append(Paragraph(f'Porcentaje de plagio general: {porcentaje_de_plagio}%', getSampleStyleSheet().get('Normal')))
    pdf_content.append(Spacer(1, 12))  # Espacio en blanco entre párrafos

    # Crear una lista con los datos de plagio para la tabla
    data = [['Oración plagiada', 'Oración original', 'Lugar donde se encontró', 'Ubicación']]
    for oracion, plagioE, porcentaje, url, ubicacion in plagio:
        # print("algo aqui")
        data.append([Paragraph(oracion, getSampleStyleSheet().get('Normal')),
                     Paragraph(plagioE, getSampleStyleSheet().get('Normal')),
                     Paragraph(url, getSampleStyleSheet().get('Normal')),
                     Paragraph(ubicacion, getSampleStyleSheet().get('Normal'))])

    # Agregar la tabla al PDF
    table_style = TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.gray),
                              ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                              ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                              ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                              ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                              ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                              ('GRID', (0, 0), (-1, -1), 1, colors.black)
                              # Controlar el flujo de contenido en las celdas
                              ])

    # Ajustar automáticamente el tamaño de las celdas
    col_widths = [142, 142, 120, 61]  # Ajustar el ancho de las columnas basado en el contenido más largo
    table = Table(data, style=table_style, colWidths=col_widths)
    pdf_content.append(table)

    # Crear el documento PDF
    nombre_archivo_pdf = path_resultado + 'Plagio ' + str(str(nombre_archivo).split(".")[0]) + '.pdf'
    pdf = SimpleDocTemplate(nombre_archivo_pdf, pagesize=letter)
    pdf.build(pdf_content)
    nombre = 'Plagio ' + str(str(nombre_archivo).split(".")[0]) + '.pdf'
    return nombre_archivo_pdf, nombre
